import csv

with open('D:\pyhton_Practice\data_file.csv') as csv_file:
    csv_reader = csv.reader(csv_file, delimiter=',')
    main_data = []
    line_count = 0
    
    for row in csv_reader:
        # if(line_count > 8):
        #     break;
        
        if(line_count != 0):
            if(len(row) != 0):
                new_dict = {}
                new_dict['name'] = row[0]
                new_dict['pan_no'] = row[2]
                new_dict['mobile_no'] = row[3]
                new_dict['current_add'] = row[5]
                new_dict['permanent_add'] = row[12]
                new_dict['current_add'] = row[5]
                main_data.append(new_dict)
               
        line_count += 1
    print('Processed {line_count} lines.')
    print(main_data)
    print(len(main_data))

# for editing doc file
from docx import Document

# document = Document('D:\pyhton_Practice\suma.docx')
# tbl = document.tables[1]
# #print(tbl.rows[0].cells[1].text)

# tbl.rows[0].cells[1].text = ""
# p = tbl.rows[0].cells[1].add_paragraph()
# p.add_run('Borrower Name:').bold = True
# p.add_run('Borrower Name')
# tbl.rows[0].cells[1].add_paragraph("")

# w = tbl.rows[0].cells[1].add_paragraph()
# w.add_run('Current Address:').bold = True
# w.add_run('Borrower Name')
# tbl.rows[0].cells[1].add_paragraph("")
# e = tbl.rows[0].cells[1].add_paragraph()
# e.add_run('Current Address:').bold = True
# e.add_run('Borrower Name')
# tbl.rows[0].cells[1].add_paragraph("")

# r = tbl.rows[0].cells[1].add_paragraph()
# r.add_run('Permanent Address:').bold = True
# r.add_run('Borrower Name')
# tbl.rows[0].cells[1].add_paragraph("")

# t = tbl.rows[0].cells[1].add_paragraph()
# t.add_run('Contact details:').bold = True
# t.add_run('Borrower Name')
# tbl.rows[0].cells[1].add_paragraph("")

# y = tbl.rows[0].cells[1].add_paragraph()
# y.add_run('PAN Card no:').bold = True
# y.add_run('Borrower Name')

# print(tbl.rows[0].cells[1].text)
# document.save('test1rty.docx')

# 